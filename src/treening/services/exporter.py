"""Portable exports for textbook-learning quiz sessions."""

from __future__ import annotations

import re
from io import BytesIO
from typing import Any


BRANCH_LABELS = {
    "question": "起点问题",
    "check": "验收",
    "followup": "追问",
    "custom": "其他",
    "correction": "其他",
}


def _branch(node: dict[str, Any]) -> str:
    value = node.get("branch_type") or "question"
    return "custom" if value == "correction" else value


def _selected_nodes(
    nodes: list[dict[str, Any]],
    scope: str,
    node_id: str | None,
) -> list[dict[str, Any]]:
    if scope not in {"tree", "path", "subtree"}:
        raise ValueError("invalid export scope")
    if not nodes:
        return []

    node_map = {node["id"]: node for node in nodes}
    target_id = node_id or nodes[-1]["id"]
    if scope == "tree":
        return nodes
    if target_id not in node_map:
        raise ValueError("export node not found")
    if scope == "path":
        result: list[dict[str, Any]] = []
        current_id: str | None = target_id
        seen: set[str] = set()
        while current_id and current_id not in seen:
            seen.add(current_id)
            current = node_map.get(current_id)
            if not current:
                break
            result.append(current)
            current_id = current.get("parent_id")
        result.reverse()
        return result

    descendants = {target_id}
    changed = True
    while changed:
        changed = False
        for node in nodes:
            if node.get("parent_id") in descendants and node["id"] not in descendants:
                descendants.add(node["id"])
                changed = True
    return [node for node in nodes if node["id"] in descendants]


def _depths(nodes: list[dict[str, Any]]) -> dict[str, int]:
    node_map = {node["id"]: node for node in nodes}
    memo: dict[str, int] = {}

    def depth(node_id: str, trail: set[str] | None = None) -> int:
        if node_id in memo:
            return memo[node_id]
        trail = trail or set()
        if node_id in trail:
            return 0
        node = node_map[node_id]
        parent_id = node.get("parent_id")
        value = depth(parent_id, trail | {node_id}) + 1 if parent_id in node_map else 0
        memo[node_id] = value
        return value

    return {node["id"]: depth(node["id"]) for node in nodes}


def _tree_lines(nodes: list[dict[str, Any]]) -> list[str]:
    depths = _depths(nodes)
    return [
        f"{'  ' * depths[node['id']]}- [{BRANCH_LABELS.get(_branch(node), '学习')}] "
        f"{node.get('role', 'unknown')}: {str(node.get('content', '')).strip()}"
        for node in nodes
    ]


def _mermaid(nodes: list[dict[str, Any]]) -> str:
    lines = ["```mermaid", "flowchart TD"]
    for index, node in enumerate(nodes, start=1):
        node_key = f"N{index}"
        label = re.sub(r"[\r\n]+", " ", str(node.get("content", ""))).strip()
        label = label.replace('"', "'")[:72] or "空节点"
        lines.append(f'  {node_key}["{label}"]')
    id_to_key = {node["id"]: f"N{index}" for index, node in enumerate(nodes, start=1)}
    for node in nodes:
        parent_id = node.get("parent_id")
        if parent_id in id_to_key:
            lines.append(f"  {id_to_key[parent_id]} -->|{BRANCH_LABELS.get(_branch(node), '学习')}| {id_to_key[node['id']]}")
    lines.append("```")
    return "\n".join(lines)


_DECON_LABELS = (
    ("contradiction", "矛盾论 · 认识拆解"),
    ("practice", "实践论 · 行动指向"),
    ("check_question", "问题 · 验收"),
    ("reflect_question", "问题 · 反思"),
    ("inspire_question", "问题 · 启发"),
)


def _deconstruction_blocks(node: dict[str, Any]) -> list[tuple[str, str]]:
    """Extract non-empty 拆解 blocks (矛盾论/实践论/三问) from a node.

    Blocks live in assistant node metadata. Empty or missing entries are
    skipped so honest "现阶段不需要行动" responses do not leave a bare label.
    """
    metadata = node.get("metadata")
    if not isinstance(metadata, dict):
        return []
    return [
        (label, str(metadata.get(key, "")).strip())
        for key, label in _DECON_LABELS
        if metadata.get(key)
    ]


def _title(session: dict[str, Any]) -> str:
    return (session.get("title") or session.get("root_question") or "未命名学习主题").strip()[:120]


def render_markdown(session: dict[str, Any], nodes: list[dict[str, Any]], scope: str) -> str:
    title = _title(session)
    lines = [
        f"# {title}",
        "",
        f"- 创建时间：{session.get('created_at', '')}",
        f"- 最近更新：{session.get('updated_at', '')}",
        f"- 导出范围：{scope}",
        f"- 节点数量：{len(nodes)}",
        "",
        "## 知识树",
        "",
        _mermaid(nodes),
        "",
        "## 节点内容",
        "",
    ]
    depths = _depths(nodes)
    for index, node in enumerate(nodes, start=1):
        branch = BRANCH_LABELS.get(_branch(node), "学习")
        role = "提问" if node.get("role") == "user" else "回答"
        lines.extend([
            f"### {index}. {role} · {branch}",
            "",
            f"> 深度 {depths[node['id']] + 1} · 节点 ID `{node['id']}`",
            "",
            str(node.get("content", "")).strip(),
            "",
        ])
        blocks = _deconstruction_blocks(node)
        if blocks:
            lines.extend(["**拆解**", ""])
            for label, text in blocks:
                lines.append(f"- **{label}**：{text}")
            lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def render_text(session: dict[str, Any], nodes: list[dict[str, Any]], scope: str) -> str:
    lines = [
        _title(session),
        "=" * len(_title(session)),
        f"创建时间：{session.get('created_at', '')}",
        f"最近更新：{session.get('updated_at', '')}",
        f"导出范围：{scope}",
        "",
        "知识树：",
        *_tree_lines(nodes),
        "",
    ]
    lines.append("")
    for node in nodes:
        blocks = _deconstruction_blocks(node)
        if not blocks:
            continue
        branch = BRANCH_LABELS.get(_branch(node), "学习")
        lines.append(f"[拆解 · {branch}]")
        for label, text in blocks:
            lines.append(f"  {label}：{text}")
        lines.append("")
    return "\n".join(lines)


def render_docx(session: dict[str, Any], nodes: list[dict[str, Any]], scope: str) -> bytes:
    from docx import Document
    from docx.shared import Pt

    document = Document()
    title = _title(session)
    document.add_heading(title, 0)
    document.add_paragraph(f"创建时间：{session.get('created_at', '')}")
    document.add_paragraph(f"最近更新：{session.get('updated_at', '')}")
    document.add_paragraph(f"导出范围：{scope} · 节点数量：{len(nodes)}")
    document.add_heading("知识树", level=1)
    for line in _tree_lines(nodes):
        paragraph = document.add_paragraph(line)
        paragraph.paragraph_format.left_indent = Pt(12)
    document.add_heading("节点内容", level=1)
    depths = _depths(nodes)
    for index, node in enumerate(nodes, start=1):
        branch = BRANCH_LABELS.get(_branch(node), "学习")
        role = "提问" if node.get("role") == "user" else "回答"
        document.add_heading(f"{index}. {role} · {branch}", level=min(9, depths[node["id"]] + 2))
        document.add_paragraph(str(node.get("content", "")).strip())
        blocks = _deconstruction_blocks(node)
        if blocks:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(12)
            run = paragraph.add_run("拆解")
            run.bold = True
            for label, text in blocks:
                item = document.add_paragraph(style="List Bullet")
                item.paragraph_format.left_indent = Pt(24)
                label_run = item.add_run(f"{label}：")
                label_run.bold = True
                item.add_run(text)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


# ═══════════════════════════════════════════════════════════════════════════
# Obsidian vault 导出：多 md + frontmatter + wikilink + MOC
# ═══════════════════════════════════════════════════════════════════════════

def _slugify(text: str | None, fallback: str = "note") -> str:
    """把标题转成安全的文件名/ID 片段（保留中文、字母、数字、连字符）。"""
    value = str(text or "").strip()
    value = re.sub(r"[\[\]#*`|]", "", value)
    value = re.sub(r"[^\w一-鿿-]+", "-", value, flags=re.UNICODE)
    value = value.strip("-")
    return value[:60] or fallback


_SCOPE_SUFFIX = {
    "tree": "",
    "path": "路径",
    "subtree": "子树",
}


def export_basename(session: dict[str, Any], scope: str = "tree") -> str:
    """统一的导出文件名（不含扩展名）。

    下载名与 vault 内部文件夹名共用同一个函数，保证「下载 .zip 解压出的文件夹」
    与「文件名本身」一致；非整树导出会追加范围后缀（路径 / 子树）以便区分。
    """
    base = _slugify(_title(session), "tree-session")
    suffix = _SCOPE_SUFFIX.get(scope or "tree", "")
    return f"{base}-{suffix}" if suffix else base


def _iso_date(value: Any) -> str:
    return str(value or "").strip()[:10]


def _yaml_str(value: Any) -> str:
    escaped = str(value or "").replace("\\", "\\\\").replace('"', '\\"')
    return f'"{escaped}"'


def _meta(node: dict[str, Any], key: str) -> str:
    metadata = node.get("metadata")
    if isinstance(metadata, dict):
        raw = metadata.get(key)
        if raw:
            return str(raw).strip()
    return ""


def _note_title(answer: dict[str, Any], question: dict[str, Any] | None) -> str:
    """笔记标题：优先用提问（最自然），其次召回摘要，最后取回答首句。"""
    if question:
        q = str(question.get("content", "")).strip()
        first = q.splitlines()[0].strip() if q else ""
        if first:
            return first[:80]
    summary = _meta(answer, "summary")
    if summary:
        return summary[:80]
    content = str(answer.get("content", "")).strip()
    first = content.splitlines()[0].strip() if content else ""
    return first[:80] or "未命名概念"


def _answer_topology(
    nodes: list[dict[str, Any]],
) -> tuple[dict[str, dict[str, Any]], dict[str, list[str]], dict[str, str], dict[str, list[str]]]:
    """构建回答节点的拓扑：节点映射、父子索引、提问映射、子回答索引。

    返回：
      node_map   id -> node
      children   父 id -> 子 id 列表（按 created_at 排序）
      question_of  回答 id -> 其提问（user）节点 id
      child_answers_of  回答 id -> 其下游回答（assistant）节点 id 列表
    """
    node_map = {node["id"]: node for node in nodes}
    children: dict[str, list[str]] = {}
    for node in nodes:
        parent_id = node.get("parent_id")
        if parent_id and parent_id in node_map:
            children.setdefault(parent_id, []).append(node["id"])
    for ids in children.values():
        ids.sort(key=lambda nid: node_map[nid].get("created_at") or "")

    question_of: dict[str, str] = {}
    child_answers_of: dict[str, list[str]] = {}
    for node in nodes:
        if node.get("role") != "assistant":
            continue
        parent = node_map.get(node.get("parent_id") or "")
        if parent and parent.get("role") == "user":
            question_of[node["id"]] = parent["id"]
            grand = parent.get("parent_id")
            if grand:
                child_answers_of.setdefault(grand, []).append(node["id"])
    return node_map, children, question_of, child_answers_of


def _assign_unique_titles(
    answer_ids: list[str],
    node_map: dict[str, dict[str, Any]],
    question_of: dict[str, str],
) -> dict[str, str]:
    seen: dict[str, int] = {}
    titles: dict[str, str] = {}
    for aid in answer_ids:
        question = node_map.get(question_of.get(aid, ""))
        base = _note_title(node_map[aid], question)
        if base in seen:
            seen[base] += 1
            unique = f"{base} ({seen[base]})"
        else:
            seen[base] = 1
            unique = base
        titles[aid] = unique
    return titles


def _concept_body(
    answer: dict[str, Any],
    question: dict[str, Any] | None,
    parent_title: str | None,
    child_titles: list[str],
) -> str:
    """生成 concept 笔记正文：先「问题 → 回答」打包问答，再拆解五字段，最后关联。

    刻意让问答对成为笔记的第一主体（可读性优先）；拆解沿用产品里已有的中文标签
    （矛盾论 / 实践论 / 三问），避免英文模板节名带来的理解成本；有内容才渲染对应
    小节，空小节不再占位，正文更干净。
    """
    qtext = str(question.get("content", "")).strip() if question else ""
    content = str(answer.get("content", "")).strip()
    contradiction = _meta(answer, "contradiction")
    practice = _meta(answer, "practice")
    check_q = _meta(answer, "check_question")
    reflect_q = _meta(answer, "reflect_question")
    inspire_q = _meta(answer, "inspire_question")

    parts: list[str] = []

    # 问答对：打包在最前，作为笔记的核心
    parts.append("## 问题")
    parts.append("")
    parts.append(qtext or "（无）")
    parts.append("")

    parts.append("## 回答")
    parts.append("")
    parts.append(content or "（无）")
    parts.append("")

    # 拆解：沿用「矛盾论 / 实践论 / 三问」的中文标签，有值才列出
    decons: list[str] = []
    if contradiction:
        decons.append(f"- **矛盾论 · 认识拆解**：{contradiction}")
    if practice:
        decons.append(f"- **实践论 · 行动指向**：{practice}")
    if check_q:
        decons.append(f"- **问题 · 验收**：{check_q}")
    if reflect_q:
        decons.append(f"- **问题 · 反思**：{reflect_q}")
    if inspire_q:
        decons.append(f"- **问题 · 启发**：{inspire_q}")
    if decons:
        parts.append("## 拆解")
        parts.append("")
        parts.extend(decons)
        parts.append("")

    # 关联：区分上级与分支，便于在 Obsidian 里顺藤摸瓜
    related: list[str] = []
    if parent_title:
        related.append(f"- ↑ 上级：[[{parent_title}]]")
    if child_titles:
        related.append("- ↓ 分支：" + " · ".join(f"[[{t}]]" for t in child_titles))
    if related:
        parts.append("## 关联")
        parts.append("")
        parts.extend(related)
        parts.append("")

    return "\n".join(parts).rstrip() + "\n"


def _has_check_child(answer_id: str, children: dict[str, list[str]], node_map: dict[str, dict[str, Any]]) -> bool:
    for cid in children.get(answer_id, []):
        child = node_map.get(cid)
        if child and child.get("role") == "user" and child.get("branch_type") == "check":
            return True
    return False


def _render_note(
    title: str,
    slug: str,
    status: str,
    created: str,
    updated: str,
    session_slug: str,
    body: str,
) -> str:
    frontmatter = [
        "---",
        f"id: {_yaml_str(slug)}",
        "type: concept",
        "level: L2",
        f"status: {status}",
        f"title: {_yaml_str(title)}",
        f"tags: [treening, {_yaml_str(session_slug)}]",
    ]
    if created:
        frontmatter.append(f"created: {created}")
    if updated:
        frontmatter.append(f"updated: {updated}")
    frontmatter.append("---")
    frontmatter.append("")
    frontmatter.append(f"# {title}")
    frontmatter.append("")
    return "\n".join(frontmatter) + body + "\n"


def _render_moc(
    session: dict[str, Any],
    root_ids: list[str],
    child_answers_of: dict[str, list[str]],
    titles: dict[str, str],
) -> str:
    title = _title(session)
    session_slug = _slugify(title, "tree")
    lines = [
        "---",
        f"id: {_yaml_str(session_slug + '-moc')}",
        "type: moc",
        f"title: {_yaml_str(title)}",
    ]
    if session.get("created_at"):
        lines.append(f"created: {_iso_date(session.get('created_at'))}")
    if session.get("updated_at"):
        lines.append(f"updated: {_iso_date(session.get('updated_at'))}")
    lines += ["---", "", f"# {title}", ""]
    root_question = str(session.get("root_question") or "").strip()
    if root_question:
        lines += [f"> 起点问题：{root_question}", ""]
    lines += ["## 知识树", ""]

    def walk(aid: str, depth: int) -> None:
        lines.append(f"{'  ' * depth}- [[{titles[aid]}]]")
        for cid in child_answers_of.get(aid, []):
            walk(cid, depth + 1)

    for rid in root_ids:
        walk(rid, 0)
    return "\n".join(lines) + "\n"


# ── Obsidian Canvas ──
# 卡片尺寸为原默认的 4 倍面积（长宽各 ×2），间距同步放大，保证排版舒展不拥挤。
_CANVAS_NODE_WIDTH = 640
_CANVAS_NODE_HEIGHT = 240
_CANVAS_H_GAP = 160
_CANVAS_V_GAP = 120

# Canvas 卡片颜色（Obsidian 只接受预设 1..6，不接受 hex）：
#   起点(根) 用紫、验收 用绿、追问 用蓝、其他 用橙。
_CANVAS_COLORS = {
    "question": "6",
    "check": "3",
    "followup": "4",
    "custom": "2",
}

# 横向 canvas 的间距：深度层之间的 x 间距、同层兄弟之间的 y 间距。
# 语义与纵向的 _CANVAS_H_GAP / _CANVAS_V_GAP 正好互换（横向时 x 是「层」、y 是「分支展开」）。
_CANVAS_H_LAYER_GAP = 160
_CANVAS_V_SIBLING_GAP = 120


def _canvas_layout(
    root_ids: list[str],
    child_answers_of: dict[str, list[str]],
) -> dict[str, tuple[float, float]]:
    """简单的 tidy-tree 布局（自上而下）。

    叶子按中序遍历从左到右各占一个槽位；内部节点的 x 取子节点 x 的中点，
    y 按深度递增。结果保证同一层不重叠、分支按左右展开。
    """
    positions: dict[str, tuple[float, float]] = {}
    depth: dict[str, int] = {}

    def set_depth(aid: str, d: int) -> None:
        depth[aid] = d
        for kid in child_answers_of.get(aid, []):
            set_depth(kid, d + 1)

    for rid in root_ids:
        set_depth(rid, 0)

    leaf_index: dict[str, int] = {}

    def collect_leaves(aid: str) -> None:
        kids = child_answers_of.get(aid, [])
        if not kids:
            leaf_index[aid] = len(leaf_index)
        else:
            for kid in kids:
                collect_leaves(kid)

    for rid in root_ids:
        collect_leaves(rid)

    def assign_x(aid: str) -> float:
        kids = child_answers_of.get(aid, [])
        if not kids:
            x = leaf_index[aid] * (_CANVAS_NODE_WIDTH + _CANVAS_H_GAP)
        else:
            x = sum(assign_x(kid) for kid in kids) / len(kids)
        positions[aid] = (x, depth[aid] * (_CANVAS_NODE_HEIGHT + _CANVAS_V_GAP))
        return x

    for rid in root_ids:
        assign_x(rid)
    return positions


def _canvas_layout_horizontal(
    root_ids: list[str],
    child_answers_of: dict[str, list[str]],
) -> dict[str, tuple[float, float]]:
    """横向河流式布局：x 随深度递增（层），兄弟子树按叶子槽位纵向堆叠（y）。

    与 _canvas_layout 互为镜像：横向屏幕下回答卡沿 x 向右生长，
    同一深度的兄弟在 y 方向错开，观感与学习空间的横向模式一致。
    """
    positions: dict[str, tuple[float, float]] = {}
    depth: dict[str, int] = {}

    def set_depth(aid: str, d: int) -> None:
        depth[aid] = d
        for kid in child_answers_of.get(aid, []):
            set_depth(kid, d + 1)

    for rid in root_ids:
        set_depth(rid, 0)

    leaf_index: dict[str, int] = {}

    def collect_leaves(aid: str) -> None:
        kids = child_answers_of.get(aid, [])
        if not kids:
            leaf_index[aid] = len(leaf_index)
        else:
            for kid in kids:
                collect_leaves(kid)

    for rid in root_ids:
        collect_leaves(rid)

    def assign_y(aid: str) -> float:
        kids = child_answers_of.get(aid, [])
        if not kids:
            y = leaf_index[aid] * (_CANVAS_NODE_HEIGHT + _CANVAS_V_SIBLING_GAP)
        else:
            y = sum(assign_y(kid) for kid in kids) / len(kids)
        positions[aid] = (depth[aid] * (_CANVAS_NODE_WIDTH + _CANVAS_H_LAYER_GAP), y)
        return y

    for rid in root_ids:
        assign_y(rid)
    return positions


def _render_canvas(
    session: dict[str, Any],
    root_ids: list[str],
    child_answers_of: dict[str, list[str]],
    slugs: dict[str, str],
    branch_types: dict[str, str],
    folder: str,
    layout: str = "vertical",
) -> str:
    """把回答树渲染成 Obsidian Canvas（.canvas JSON）。

    每个回答一张 ``file`` 节点（指向 vault 里的 md），父子用带箭头的边连接；
    用会话标题作为一张 ``text`` 根节点锚定整张画布。
    layout 决定画布方向：
    - vertical   自上而下（默认）：标题在根上方，边 bottom→top
    - horizontal 横向河流：x 随深度递增、兄弟纵向堆叠，标题在根左侧，边 right→left
    """
    import json as _json

    horizontal = layout == "horizontal"
    positions = (
        _canvas_layout_horizontal(root_ids, child_answers_of)
        if horizontal
        else _canvas_layout(root_ids, child_answers_of)
    )
    nodes: list[dict[str, Any]] = []
    edges: list[dict[str, Any]] = []

    _TITLE_WIDTH = 480
    _TITLE_HEIGHT = 80

    if horizontal:
        # 标题卡居左、垂直居中对齐根回答；节点层从标题右侧开始，x = 标题宽 + 层间距
        _LEFT_OFFSET = _TITLE_WIDTH + _CANVAS_H_LAYER_GAP
        anchor_y = positions[root_ids[0]][1] if root_ids else 0.0
        nodes.append({
            "id": "canvas-title",
            "type": "text",
            "text": _title(session),
            "x": 0.0,
            "y": round(anchor_y + (_CANVAS_NODE_HEIGHT - _TITLE_HEIGHT) / 2, 1),
            "width": _TITLE_WIDTH,
            "height": _TITLE_HEIGHT,
            "color": "1",
        })
        for aid, (x, y) in positions.items():
            nodes.append({
                "id": aid,
                "type": "file",
                "file": f"{folder}/{slugs[aid]}.md",
                "x": round(_LEFT_OFFSET + x, 1),
                "y": round(y, 1),
                "width": _CANVAS_NODE_WIDTH,
                "height": _CANVAS_NODE_HEIGHT,
                "color": _CANVAS_COLORS.get(branch_types.get(aid, "custom"), "2"),
            })
    else:
        # 标题卡片：居中锚定在根节点正上方，底部引一条竖线连到根（比斜线更规整）。
        _TOP_OFFSET = _TITLE_HEIGHT + _CANVAS_V_GAP
        anchor_x = positions[root_ids[0]][0] if root_ids else 0.0
        nodes.append({
            "id": "canvas-title",
            "type": "text",
            "text": _title(session),
            "x": round(anchor_x + (_CANVAS_NODE_WIDTH - _TITLE_WIDTH) / 2, 1),
            "y": 0.0,
            "width": _TITLE_WIDTH,
            "height": _TITLE_HEIGHT,
            "color": "1",
        })
        for aid, (x, y) in positions.items():
            nodes.append({
                "id": aid,
                "type": "file",
                "file": f"{folder}/{slugs[aid]}.md",
                "x": round(x, 1),
                "y": round(_TOP_OFFSET + y, 1),
                "width": _CANVAS_NODE_WIDTH,
                "height": _CANVAS_NODE_HEIGHT,
                "color": _CANVAS_COLORS.get(branch_types.get(aid, "custom"), "2"),
            })

    title_side = ("right", "left") if horizontal else ("bottom", "top")
    for rid in root_ids:
        edges.append({
            "id": f"title-{rid}",
            "fromNode": "canvas-title",
            "fromSide": title_side[0],
            "fromEnd": "none",
            "toNode": rid,
            "toSide": title_side[1],
            "toEnd": "arrow",
        })

    for aid in positions:
        for kid in child_answers_of.get(aid, []):
            edges.append({
                "id": f"{aid}-{kid}",
                "fromNode": aid,
                "fromSide": title_side[0],
                "fromEnd": "none",
                "toNode": kid,
                "toSide": title_side[1],
                "toEnd": "arrow",
                "label": BRANCH_LABELS.get(branch_types.get(kid, "custom"), "其他"),
            })

    return _json.dumps({"nodes": nodes, "edges": edges}, ensure_ascii=False, indent=2)


def render_vault(
    session: dict[str, Any],
    nodes: list[dict[str, Any]],
    scope: str = "tree",
    layout: str = "vertical",
) -> bytes:
    """整棵树导出为 Obsidian vault 的 zip：多 md + frontmatter + wikilink + MOC。

    ``nodes`` 应已是按 scope/node_id 过滤后的节点列表（由 render_export 负责过滤）；
    ``scope`` 只用于确定 zip 内的文件夹名（与下载名共用 export_basename）。
    ``layout`` 控制 .canvas 画布方向（vertical / horizontal）。
    """
    import zipfile

    node_map, children, question_of, child_answers_of = _answer_topology(nodes)
    answer_ids = [n["id"] for n in nodes if n.get("role") == "assistant"]

    titles = _assign_unique_titles(answer_ids, node_map, question_of)
    session_slug = export_basename(session, scope)

    # 每个回答节点是一个「概念」笔记；根回答的父回答为 None
    parent_answer_of: dict[str, str | None] = {}
    for aid in answer_ids:
        question_id = question_of.get(aid)
        question = node_map.get(question_id or "")
        parent_answer_of[aid] = question.get("parent_id") if question else None

    files: dict[str, str] = {}
    slugs: dict[str, str] = {}
    for aid in answer_ids:
        question_id = question_of.get(aid)
        question = node_map.get(question_id or "")
        parent_title = titles.get(parent_answer_of[aid]) if parent_answer_of[aid] else None
        child_titles = [
            titles[cid] for cid in child_answers_of.get(aid, []) if cid in titles
        ]
        status = "verified" if _has_check_child(aid, children, node_map) else "draft"
        title = titles[aid]
        slug = _slugify(title, "concept")
        slugs[aid] = slug
        created = _iso_date(node_map[aid].get("created_at") or session.get("created_at"))
        updated = _iso_date(session.get("updated_at"))
        body = _concept_body(node_map[aid], question, parent_title, child_titles)
        files[f"{slug}.md"] = _render_note(
            title, slug, status, created, updated, session_slug, body
        )

    root_ids = [aid for aid in answer_ids if not parent_answer_of.get(aid)]
    branch_types = {aid: _branch(node_map[aid]) for aid in answer_ids}
    files["MOC.md"] = _render_moc(session, root_ids, child_answers_of, titles)
    files[f"{session_slug}.canvas"] = _render_canvas(
        session, root_ids, child_answers_of, slugs, branch_types, session_slug, layout
    )

    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, content in files.items():
            archive.writestr(f"{session_slug}/{name}", content.encode("utf-8"))
    return buffer.getvalue()


def render_export(
    session: dict[str, Any],
    nodes: list[dict[str, Any]],
    fmt: str,
    scope: str,
    node_id: str | None = None,
    layout: str = "vertical",
) -> tuple[bytes, str, str]:
    selected = _selected_nodes(nodes, scope, node_id)
    if fmt == "md":
        return render_markdown(session, selected, scope).encode("utf-8"), "text/markdown; charset=utf-8", "md"
    if fmt == "txt":
        return render_text(session, selected, scope).encode("utf-8"), "text/plain; charset=utf-8", "txt"
    if fmt == "docx":
        return render_docx(session, selected, scope), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"
    if fmt in ("vault", "obsidian"):
        return render_vault(session, selected, scope, layout), "application/zip", "zip"
    raise ValueError("unsupported export format")
